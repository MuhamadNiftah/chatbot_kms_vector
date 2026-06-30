import sys
sys.stdout.reconfigure(encoding='utf-8')
import os
import json
import datetime
import csv
import threading
from typing import List, Optional

# File processing libraries
try:
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image
    import pypdf
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    print("⚠️ Warning: pdf2image, pytesseract, or pypdf not found. PDF OCR and Image support may be limited.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ Warning: python-docx not found. Word support disabled.")

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("⚠️ Warning: openpyxl not found. Excel support disabled.")

import mysql.connector
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ==============================
# CONFIG
# ==============================
folder_path = "documents"
vectorstore_dir = "vectorstore"

# Tesseract & Poppler paths setup (Windows only)
if os.name == 'nt':
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    poppler_path = r"C:\poppler\poppler-25.12.0\Library\bin"
else:
    poppler_path = None


# ==============================
# KONEKSI MYSQL
# ==============================
def get_db_connection():
    return mysql.connector.connect(
        host="localhost",
        user="root",
        password="",
        database="ai_chat"
    )

# ==============================
# EMBEDDINGS
# ==============================
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

# ==============================
# FILE PROCESSORS
# ==============================

def process_txt_file(file_path: str) -> List[Document]:
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        if not content.strip():
            return []
        filename = os.path.basename(file_path)
        return [Document(
            page_content=content,
            metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "txt", "page": 1}
        )]
    except Exception as e:
        print(f"❌ Error TXT {file_path}: {e}")
        return []

def process_csv_file(file_path: str) -> List[Document]:
    try:
        content_lines = []
        with open(file_path, 'r', encoding='utf-8') as f:
            csv_reader = csv.DictReader(f)
            for row in csv_reader:
                content_lines.append(", ".join([f"{k}: {v}" for k, v in row.items()]))
        content = "\n".join(content_lines)
        if not content.strip():
            return []
        filename = os.path.basename(file_path)
        return [Document(
            page_content=content,
            metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "csv", "page": 1}
        )]
    except Exception as e:
        print(f"❌ Error CSV {file_path}: {e}")
        return []

def process_docx_file(file_path: str) -> List[Document]:
    if not DOCX_AVAILABLE:
        print(f"⚠️ python-docx not available for {file_path}")
        return []
    try:
        doc_obj = DocxDocument(file_path)
        paragraphs = [para.text for para in doc_obj.paragraphs if para.text.strip()]
        content = "\n".join(paragraphs)
        if not content.strip():
            return []
        filename = os.path.basename(file_path)
        return [Document(
            page_content=content,
            metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "docx", "page": 1}
        )]
    except Exception as e:
        print(f"❌ Error DOCX {file_path}: {e}")
        return []

def process_excel_file(file_path: str) -> List[Document]:
    if not EXCEL_AVAILABLE:
        print(f"⚠️ openpyxl not available for {file_path}")
        return []
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        content_lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            content_lines.append(f"Sheet: {sheet_name}")
            for row in ws.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                content_lines.append(" | ".join(row_data))
        content = "\n".join(content_lines)
        if not content.strip():
            return []
        filename = os.path.basename(file_path)
        return [Document(
            page_content=content,
            metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "excel", "page": 1}
        )]
    except Exception as e:
        print(f"❌ Error EXCEL {file_path}: {e}")
        return []

def process_pdf_file(file_path: str) -> List[Document]:
    try:
        filename = os.path.basename(file_path)
        documents = []
        
        # 1. Try text extraction first
        try:
            reader = pypdf.PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "pdf", "page": i+1}
                    ))
        except Exception as e:
            print(f"⚠️ pypdf failed for {filename}: {e}")

        # 2. If no text extracted, fallback to OCR
        if not documents and TESSERACT_AVAILABLE:
            print(f"🔍 No text found in {filename}, attempting OCR...")
            images = convert_from_path(file_path, poppler_path=poppler_path, dpi=300)
            for i, img in enumerate(images):
                text = pytesseract.image_to_string(img, lang="ind+eng")
                if text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "pdf_ocr", "page": i+1}
                    ))
        
        return documents
    except Exception as e:
        print(f"❌ Error PDF {file_path}: {e}")
        return []

def process_image_file(file_path: str) -> List[Document]:
    if not TESSERACT_AVAILABLE:
        print(f"⚠️ OCR not available for {file_path}")
        return []
    try:
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang="ind+eng")
        if not text.strip():
            return []
        filename = os.path.basename(file_path)
        return [Document(
            page_content=text,
            metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "image", "page": 1}
        )]
    except Exception as e:
        print(f"❌ Error Image {file_path}: {e}")
        return []

def process_any_file(file_path: str) -> List[Document]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.txt':
        return process_txt_file(file_path)
    elif ext == '.pdf':
        return process_pdf_file(file_path)
    elif ext in ['.docx', '.doc']:
        return process_docx_file(file_path)
    elif ext in ['.xlsx', '.xls']:
        return process_excel_file(file_path)
    elif ext == '.csv':
        return process_csv_file(file_path)
    elif ext in ['.jpg', '.jpeg', '.png', '.bmp']:
        return process_image_file(file_path)
    else:
        print(f"⚠️ Unsupported format: {ext} for {file_path}")
        return []

# ==============================
# MAIN INGESTION LOGIC
# ==============================

def main():
    print("📂 Scanning folder documents...")
    all_documents = []
    
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        print(f"📁 Created folder: {folder_path}")
        return

    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):
            print(f"📄 Processing: {filename}")
            docs = process_any_file(file_path)
            if docs:
                all_documents.extend(docs)
                print(f"  ✅ Extracted {len(docs)} pages/parts")
            else:
                print(f"  ⚠️ No content extracted from {filename}")

    if not all_documents:
        print("❌ No valid documents found to ingest.")
        return

    print(f"Total documents collected: {len(all_documents)}")

    # ==============================
    # TEXT SPLIT
    # ==============================
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=700,
        chunk_overlap=150
    )

    split_docs = text_splitter.split_documents(all_documents)
    print(f"Total chunks created: {len(split_docs)}")

    # ==============================
    # SIMPAN KE MYSQL + FAISS
    # ==============================
    conn = get_db_connection()
    cursor = conn.cursor()
    now = datetime.datetime.now()
    langchain_docs = []

    print("💾 Saving to MySQL and preparing FAISS...")

    for doc in split_docs:
        content  = doc.page_content
        category = doc.metadata.get("category", "Umum")
        source   = doc.metadata.get("source", "")
        page     = doc.metadata.get("page", 0)

        # 1. Simpan ke knowledge_base
        cursor.execute("""
            INSERT INTO knowledge_base (content, category, created_at, updated_at)
            VALUES (%s, %s, %s, %s)
        """, (content, category, now, now))
        conn.commit()

        document_id = cursor.lastrowid

        # 2. Buat embedding
        vector = embeddings.embed_query(content)

        # 3. Simpan ke vector_database
        cursor.execute("""
            INSERT INTO vector_database (document_id, embedding)
            VALUES (%s, %s)
        """, (document_id, json.dumps(vector)))
        conn.commit()

        # 4. Prepare for FAISS
        doc.metadata["document_id"] = document_id
        langchain_docs.append(doc)

    # ==============================
    # VECTORSTORE FAISS
    # ==============================
    print("🔍 Updating FAISS vectorstore...")
    faiss_path = os.path.join(vectorstore_dir)
    
    if os.path.exists(os.path.join(faiss_path, "index.faiss")):
        db = FAISS.load_local(
            faiss_path,
            embeddings,
            allow_dangerous_deserialization=True
        )
        db.add_documents(langchain_docs)
        print("➕ Added to existing FAISS index")
    else:
        if not os.path.exists(vectorstore_dir):
            os.makedirs(vectorstore_dir)
        db = FAISS.from_documents(langchain_docs, embeddings)
        print("🆕 Created new FAISS index")

    db.save_local(faiss_path)

    cursor.close()
    conn.close()

    print("✅ Done! Data stored in MySQL & FAISS")

if __name__ == "__main__":
    main()
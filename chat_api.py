# -*- coding: utf-8 -*-
import sys
import io
# Fix encoding untuk Windows (cp1252 tidak support emoji)
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from fastapi import FastAPI, UploadFile, File, BackgroundTasks
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from dotenv import load_dotenv
from fastapi.middleware.cors import CORSMiddleware
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import mysql.connector
import os
import requests
import threading
import time
import json
import csv
from datetime import datetime
from typing import Optional

# File handling imports
try:
    from pdf2image import convert_from_path
    import pytesseract
    import pypdf
    TESSERACT_AVAILABLE = True
except:
    TESSERACT_AVAILABLE = False
    print("⚠️ pytesseract/pdf2image/pypdf not available - PDF support might be limited")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available - DOCX support disabled")

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except:
    EXCEL_AVAILABLE = False
    print("⚠️ openpyxl not available - EXCEL support disabled")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(BASE_DIR, ".env"))

app = FastAPI()

# ==============================
# GLOBAL VARIABLES
# ==============================
db_lock = threading.Lock()
processed_files = set()

# ==============================
# CORS
# ==============================
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==============================
# SCHEMA
# ==============================
class ChatRequest(BaseModel):
    user_id: Optional[int] = None
    conversation_id: Optional[str] = None
    message: str
    model: Optional[str] = "deepseek-chat"

class AddKnowledgeRequest(BaseModel):
    title: str
    content: str
    category: str = "Umum"

class UserRequest(BaseModel):
    name: str

class AddLLMRequest(BaseModel):
    model_name: str

# ✅ Model untuk endpoint /api/save-chat — disesuaikan dengan frontend
class SaveChatRequest(BaseModel):
    conversation_id: Optional[str] = None
    question: str
    answer: str

# ==============================
# TEXT SPLITTER
# ==============================
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=2000,
    chunk_overlap=200
)

# ==============================
# DATABASE CONNECTION
# ==============================
def get_db():
    return mysql.connector.connect(
        host=os.getenv("DB_HOST", "localhost"),
        user=os.getenv("DB_USER", "root"),
        password=os.getenv("DB_PASSWORD", ""),
        database=os.getenv("DB_NAME", "ai_chat")
    )

# ==============================
# LOAD EMBEDDING
# ==============================
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

# ==============================
# LOAD VECTORSTORE
# ==============================
try:
    vectorstore_path = os.path.join(BASE_DIR, "vectorstore")
    db = FAISS.load_local(
        vectorstore_path,
        embeddings,
        allow_dangerous_deserialization=True
    )
    print(f"✅ Vectorstore berhasil dimuat dari {vectorstore_path}")
except Exception as e:
    db = None
    print("⚠️ Vectorstore gagal dimuat:", e)


# ==============================
# SIMPAN KE knowledge_base & vector_database
# ==============================
def save_to_knowledge_base(content: str, category: str, source: str) -> int:
    try:
        conn = get_db()
        cursor = conn.cursor()
        now = datetime.now()

        cursor.execute("""
            INSERT INTO knowledge_base (content, category, created_at, updated_at)
            VALUES (%s, %s, %s, %s)
        """, (content, category, now, now))
        conn.commit()
        document_id = cursor.lastrowid

        vector = embeddings.embed_query(content)

        cursor.execute("""
            INSERT INTO vector_database (document_id, embedding)
            VALUES (%s, %s)
        """, (document_id, json.dumps(vector)))
        conn.commit()

        cursor.close()
        conn.close()
        return document_id

    except Exception as e:
        print(f"❌ Gagal simpan ke knowledge_base: {e}")
        return -1


# ==============================
# TAMBAH DOKUMEN KE VECTORSTORE
# ==============================
def add_documents_to_vectorstore(new_docs, category="Umum"):
    global db

    with db_lock:
        try:
            if len(new_docs) == 0:
                return {"status": "error", "message": "Tidak ada dokumen valid"}

            split_docs = text_splitter.split_documents(new_docs)

            enriched_docs = []
            for doc in split_docs:
                cat = doc.metadata.get("category", category)
                source = doc.metadata.get("source", "unknown")
                document_id = save_to_knowledge_base(doc.page_content, cat, source)
                doc.metadata["document_id"] = document_id
                enriched_docs.append(doc)

            if db is None:
                db = FAISS.from_documents(enriched_docs, embeddings)
                print(f"✅ Vectorstore baru dibuat dengan {len(enriched_docs)} chunks")
            else:
                db.add_documents(enriched_docs)
                print(f"✅ {len(enriched_docs)} chunks ditambahkan ke vectorstore")

            db.save_local(os.path.join(BASE_DIR, "vectorstore"))

            return {
                "status": "success",
                "message": f"{len(enriched_docs)} dokumen berhasil diproses",
                "chunks": len(enriched_docs)
            }
        except Exception as e:
            print(f"❌ Error menambah dokumen: {e}")
            return {"status": "error", "message": str(e)}


# ==============================
# PROSES FILE
# ==============================
def process_txt_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        if not content.strip():
            return None
        filename = os.path.basename(file_path)
        return Document(
            page_content=content,
            metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "txt"}
        )
    except Exception as e:
        print(f"❌ Error TXT {file_path}: {e}")
        return None


def process_csv_file(file_path):
    try:
        content_lines = []
        with open(file_path, 'r', encoding='utf-8') as f:
            csv_reader = csv.DictReader(f)
            for row in csv_reader:
                content_lines.append(", ".join([f"{k}: {v}" for k, v in row.items()]))
        content = "\n".join(content_lines)
        if not content.strip():
            return None
        filename = os.path.basename(file_path)
        return Document(
            page_content=content,
            metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "csv"}
        )
    except Exception as e:
        print(f"❌ Error CSV {file_path}: {e}")
        return None


def process_docx_file(file_path):
    if not DOCX_AVAILABLE:
        return None
    try:
        doc_obj = DocxDocument(file_path)
        paragraphs = [para.text for para in doc_obj.paragraphs if para.text.strip()]
        content = "\n".join(paragraphs)
        if not content.strip():
            return None
        filename = os.path.basename(file_path)
        return Document(
            page_content=content,
            metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "docx"}
        )
    except Exception as e:
        print(f"❌ Error DOCX {file_path}: {e}")
        return None


def process_excel_file(file_path):
    if not EXCEL_AVAILABLE:
        return None
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        content_lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            content_lines.append(f"Sheet: {sheet_name}")
            for row in ws.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                content_lines.append(" | ".join(row_data))
        content = "\n".join(content_lines)
        if not content.strip():
            return None
        filename = os.path.basename(file_path)
        return Document(
            page_content=content,
            metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "excel"}
        )
    except Exception as e:
        print(f"❌ Error EXCEL {file_path}: {e}")
        return None


def process_pdf_file(file_path):
    try:
        filename = os.path.basename(file_path)
        all_text = []
        
        # 1. Coba extraksi teks langsung (Fast & Accurate for digital PDF)
        try:
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    all_text.append(text)
        except Exception as e:
            print(f"⚠️ pypdf failed for {filename}: {e}")

        # 2. Jika teks sangat sedikit, coba OCR (Slow fallback for scanned PDF)
        if len("\n".join(all_text).strip()) < 100 and TESSERACT_AVAILABLE:
            print(f"🔍 Teks sedikit, mencoba OCR untuk {filename}...")
            if os.name == 'nt':
                pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
                poppler_path = r"C:\poppler\poppler-25.12.0\Library\bin"
            else:
                poppler_path = None
            images = convert_from_path(file_path, poppler_path=poppler_path, dpi=300)
            all_text = [] # Reset untuk OCR
            for img in images:
                text = pytesseract.image_to_string(img, lang="ind+eng")
                if text.strip():
                    all_text.append(text)
        
        content = "\n".join(all_text)
        if not content.strip():
            return None
            
        return Document(
            page_content=content,
            metadata={"source": filename, "category": os.path.splitext(filename)[0], "type": "pdf"}
        )
    except Exception as e:
        print(f"❌ Error PDF {file_path}: {e}")
        return None


def process_image_file(file_path):
    if not TESSERACT_AVAILABLE:
        print(f"⚠️ OCR tidak tersedia untuk {file_path}")
        return None
    try:
        from PIL import Image
        img = Image.open(file_path)
        # Gunakan Tesseract untuk extract text dari gambar
        text = pytesseract.image_to_string(img, lang="ind+eng")
        if not text.strip():
            print(f"⚠️ Gambar {os.path.basename(file_path)} tidak mengandung teks yang terdeteksi.")
            return None
        
        filename = os.path.basename(file_path)
        return Document(
            page_content=text,
            metadata={
                "source": filename, 
                "category": os.path.splitext(filename)[0], 
                "type": "image",
                "processed_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
        )
    except Exception as e:
        print(f"❌ Error Image {file_path}: {e}")
        return None


def process_file(file_path):
    file_ext = os.path.splitext(file_path)[1].lower()
    print(f"📄 Memproses {os.path.basename(file_path)} ({file_ext})...")
    
    if file_ext == '.txt':
        return process_txt_file(file_path)
    elif file_ext == '.csv':
        return process_csv_file(file_path)
    elif file_ext in ['.docx', '.doc']:
        return process_docx_file(file_path)
    elif file_ext in ['.xlsx', '.xls']:
        return process_excel_file(file_path)
    elif file_ext == '.pdf':
        return process_pdf_file(file_path)
    elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp']:
        return process_image_file(file_path)
    else:
        print(f"⚠️ Tipe file tidak didukung: {file_ext}")
        return None


# ==============================
# FILE WATCHER
# ==============================
def scan_and_process_new_files():
    global processed_files
    folder_path = os.path.join(BASE_DIR, "documents")
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        return
    supported_extensions = ['.txt', '.pdf', '.docx', '.doc', '.xlsx', '.xls', '.csv', '.jpg', '.jpeg', '.png', '.bmp']
    try:
        new_docs = []
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            if not os.path.isfile(file_path):
                continue
            if file_path in processed_files:
                continue
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext not in supported_extensions:
                continue
            doc = process_file(file_path)
            if doc:
                new_docs.append(doc)
                processed_files.add(file_path)
        if new_docs:
            result = add_documents_to_vectorstore(new_docs)
            print(f"📊 Auto-ingest: {result}")
            try:
                conn = get_db()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO document_logs (action, details, timestamp) VALUES (%s, %s, NOW())",
                    ("auto_ingest", json.dumps(result))
                )
                conn.commit()
                conn.close()
            except Exception as e:
                print(f"⚠️ Gagal log ke database: {e}")
    except Exception as e:
        print(f"❌ Error scanning folder: {e}")


def file_watcher():
    print("🔍 File watcher dimulai...")
    while True:
        try:
            scan_and_process_new_files()
            time.sleep(5)
        except Exception as e:
            print(f"❌ Error di file watcher: {e}")
            time.sleep(5)


watcher_thread = threading.Thread(target=file_watcher, daemon=True)
watcher_thread.start()
print("🚀 File watcher thread started")


# ==============================
# ENDPOINT: Root Info
# ==============================
@app.get("/")
def root():
    return {
        "name": "Chatbot KMS Vector",
        "version": "3.0",
        "status": "active",
        "supported_formats": [".txt", ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".jpg", ".jpeg", ".png", ".bmp"],
        "endpoints": {
            "chat":           "POST /chat",
            "save_chat":      "POST /api/save-chat",
            "history":        "GET  /history/{user_id}",
            "create_user":    "POST /user",
            "add_knowledge":  "POST /add-knowledge",
            "upload":         "POST /upload-document",
            "refresh":        "POST /refresh-vectorstore",
            "status":         "GET  /vectorstore-status",
            "monitor":        "GET  /admin/chat-monitor",
            "knowledge_list": "GET  /admin/knowledge-base",
            "llm_list":       "GET  /llm",
            "add_llm":        "POST /llm",
        }
    }


# ==============================
# ENDPOINT: User
# ==============================
@app.post("/user")
def create_user(req: UserRequest):
    try:
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute("INSERT INTO user (name) VALUES (%s)", (req.name,))
        conn.commit()
        user_id = cursor.lastrowid
        cursor.close()
        conn.close()
        return {"status": "success", "user_id": user_id, "name": req.name}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.get("/user/{user_id}")
def get_user(user_id: int):
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM user WHERE user_id = %s", (user_id,))
        user = cursor.fetchone()
        cursor.close()
        conn.close()
        if not user:
            return {"status": "error", "message": "User tidak ditemukan"}
        return {"status": "success", "user": user}
    except Exception as e:
        return {"status": "error", "message": str(e)}


# ==============================
# ENDPOINT: Chat (RAG + DeepSeek)
# ==============================
@app.post("/chat")
async def chat(req: ChatRequest):
    api_key = os.getenv("DEEPSEEK_API_KEY")

    if not api_key:
        return {"reply": "API key tidak ditemukan"}

    with db_lock:
        current_db = db

    if current_db is None:
        return {"reply": "Vectorstore belum dibuat. Silakan upload dokumen terlebih dahulu."}

    # Gunakan conversation_id dari request jika ada
    conversation_id = req.conversation_id

    try:
        # Menggunakan top 30 hasil semantic search untuk konteks lebih luas
        docs = current_db.similarity_search(req.message, k=30)

        if len(docs) == 0:
            context = "Tidak ada informasi relevan dalam dokumen."
        else:
            context_parts = []
            for doc in docs:
                source = doc.metadata.get("source", "unknown")
                page   = doc.metadata.get("page", "-")
                context_parts.append(f"[Sumber: {source}, Halaman: {page}]\n{doc.page_content}")
            context = "\n\n".join(context_parts)

        prompt = f"""Anda adalah AI Knowledge Management System perusahaan.
Tugas Anda adalah menjawab pertanyaan pengguna berdasarkan dokumen perusahaan.

ATURAN MENJAWAB (WAJIB DIIKUTI):
- Jawaban harus diutamakan dari konteks dokumen di bawah ini.
- Jika jawaban benar-benar tidak ditemukan di konteks, jawab: \"Informasi tidak ditemukan dalam dokumen.\"
- Gunakan informasi dari konteks untuk menyusun jawaban yang lengkap dan akurat.
- Gunakan bahasa Indonesia yang jelas dan informatif.
- Buat jawaban rapi dan mudah dibaca.
- Jika memungkinkan, gunakan poin-poin.

====================
KONTEKS DOKUMEN
====================
{context}

====================
PERTANYAAN USER
====================
{req.message}

====================
JAWABAN
===================="""

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        data = {
            "model": req.model if req.model else "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.1
        }
        try:
            response = requests.post(
                "https://api.deepseek.com/chat/completions",
                headers=headers,
                json=data,
                timeout=60
            )
        except requests.exceptions.RequestException as e:
            return {"reply": "Maaf, server tidak dapat menghubungi layanan AI. Silakan coba beberapa saat lagi."}

        if response.status_code != 200:
            return {"reply": f"Error API {response.status_code}: {response.text}"}

        reply = response.json()["choices"][0]["message"]["content"]

        # Selalu simpan ke database apapun kondisinya
        try:
            conn = get_db()
            cursor = conn.cursor()

            # Tentukan sumber jika ada (ambil dari docs metadata)
            sources = []
            if len(docs) > 0:
                for doc in docs[:3]:
                    src = doc.metadata.get("source")
                    if src and src not in sources:
                        sources.append(src)
            sumber_str = ", ".join(sources) if sources else ""

            # Simpan ke chat_history menggunakan kolom yang sesuai database
            cursor.execute("""
                INSERT INTO chat_history (pertanyaan, jawaban, sumber, created_at, updated_at)
                VALUES (%s, %s, %s, %s, %s)
            """, (req.message, reply, sumber_str, datetime.now(), datetime.now()))

            conn.commit()
            cursor.close()
            conn.close()
            print(f"✅ Chat tersimpan ke DB — pertanyaan={req.message[:30]}...")

        except Exception as e:
            print(f"⚠️ DB error saat simpan chat: {e}")

        return {
            "reply": reply,
            "conversation_id": conversation_id
        }

    except Exception as e:
        return {"reply": f"Error: {str(e)}"}


# ==============================
# ✅ ENDPOINT BARU: Save Chat dari Frontend
# ==============================
@app.post("/api/save-chat")
async def save_chat(req: SaveChatRequest):
    try:
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO chat_history (pertanyaan, jawaban, sumber, created_at, updated_at)
            VALUES (%s, %s, %s, %s, %s)
        """, (req.question, req.answer, "", datetime.now(), datetime.now()))
        conn.commit()
        history_id = cursor.lastrowid
        cursor.close()
        conn.close()
        print(f"✅ Chat tersimpan via /api/save-chat: history_id={history_id}")
        return {"status": "ok", "history_id": history_id}
    except Exception as e:
        print(f"❌ Gagal simpan chat: {e}")
        return {"status": "error", "message": str(e)}


# ==============================
# ENDPOINT: Riwayat Chat
# ==============================
@app.get("/history/{user_id}")
def get_history(user_id: int):
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("""
            SELECT id as history_id, pertanyaan as question, jawaban as answer, created_at as timestamp
            FROM chat_history
            ORDER BY created_at ASC
        """)
        history = cursor.fetchall()
        cursor.close()
        conn.close()
        return {"user_id": user_id, "history": history}
    except Exception as e:
        return {"status": "error", "message": str(e)}


# ==============================
# ENDPOINT: Upload Dokumen
# ==============================
@app.post("/upload-document")
async def upload_document(file: UploadFile = File(...)):
    try:
        supported_extensions = ['.txt', '.pdf', '.docx', '.doc', '.xlsx', '.xls', '.csv', '.jpg', '.jpeg', '.png', '.bmp']
        file_ext = os.path.splitext(file.filename)[1].lower()

        if file_ext not in supported_extensions:
            return {"status": "error", "message": f"Format tidak didukung. Gunakan: {', '.join(supported_extensions)}"}

        doc_dir = os.path.join(BASE_DIR, "documents")
        os.makedirs(doc_dir, exist_ok=True)
        file_path = os.path.join(doc_dir, file.filename)

        with open(file_path, 'wb') as f:
            content = await file.read()
            f.write(content)

        print(f"📤 File diupload: {file.filename}")

        doc = process_file(file_path)
        if doc:
            processed_files.add(file_path)
            result = add_documents_to_vectorstore([doc])

            try:
                conn = get_db()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO document_logs (action, details, timestamp) VALUES (%s, %s, NOW())",
                    ("upload_file", json.dumps({"filename": file.filename, "type": file_ext, "status": result["status"]}))
                )
                conn.commit()
                conn.close()
            except Exception as e:
                print(f"⚠️ Gagal log: {e}")

            result["file_path"] = file_path
            return result
        else:
            return {"status": "error", "message": "File kosong atau tidak valid"}

    except Exception as e:
        return {"status": "error", "message": str(e)}


# ==============================
# ENDPOINT: Add Knowledge Manual
# ==============================
@app.post("/add-knowledge")
async def add_knowledge(req: AddKnowledgeRequest):
    try:
        if not req.content.strip():
            return {"status": "error", "message": "Konten tidak boleh kosong"}

        doc = Document(
            page_content=req.content,
            metadata={
                "source": req.title,
                "category": req.category,
                "type": "manual_add"
            }
        )
        result = add_documents_to_vectorstore([doc], category=req.category)

        try:
            conn = get_db()
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO document_logs (action, details, timestamp) VALUES (%s, %s, NOW())",
                ("manual_add", json.dumps({"title": req.title, "category": req.category, "status": result["status"]}))
            )
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"⚠️ Gagal log: {e}")

        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


# ==============================
# ENDPOINT: Vectorstore Status
# ==============================
@app.get("/vectorstore-status")
async def vectorstore_status():
    try:
        if db is None:
            return {"status": "empty", "message": "Vectorstore belum dibuat"}
        index_size = db.index.ntotal if hasattr(db, 'index') else 0
        return {
            "status": "active",
            "documents": index_size,
            "processed_files": len(processed_files),
            "message": f"Vectorstore aktif dengan {index_size} dokumen"
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.post("/refresh-vectorstore")
async def refresh_vectorstore():
    try:
        global processed_files
        processed_files.clear()
        scan_and_process_new_files()
        return {"status": "success", "message": "Vectorstore direfresh"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


# ==============================
# ENDPOINT: LLM
# ==============================
@app.get("/llm")
def get_llm():
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM llm")
        data = cursor.fetchall()
        cursor.close()
        conn.close()
        return {"models": data}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.post("/llm")
def add_llm(req: AddLLMRequest):
    try:
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute("INSERT INTO llm (model_name) VALUES (%s)", (req.model_name,))
        conn.commit()
        llm_id = cursor.lastrowid
        cursor.close()
        conn.close()
        return {"status": "success", "llm_id": llm_id, "model_name": req.model_name}
    except Exception as e:
        return {"status": "error", "message": str(e)}


# ==============================
# ENDPOINT: Admin Monitor
# ==============================
@app.get("/admin/chat-monitor")
def monitor_chat():
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("""
            SELECT
                ch.history_id,
                u.name AS user_name,
                cs.user_id,
                ch.question,
                ch.answer,
                ch.timestamp
            FROM chat_history ch
            JOIN chatbot_system cs ON ch.conversation_id = cs.conversation_id
            JOIN user u ON cs.user_id = u.user_id
            ORDER BY ch.timestamp DESC
        """)
        data = cursor.fetchall()
        cursor.close()
        conn.close()
        return data
    except Exception as e:
        return {"error": str(e)}


@app.get("/admin/knowledge-base")
def get_knowledge_base():
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("""
            SELECT document_id, category, created_at, updated_at,
                   LEFT(content, 200) AS content_preview
            FROM knowledge_base
            ORDER BY created_at DESC
        """)
        data = cursor.fetchall()
        cursor.close()
        conn.close()
        return {"total": len(data), "knowledge_base": data}
    except Exception as e:
        return {"status": "error", "message": str(e)}


# ==============================
# RUN: uvicorn main:app --reload --port 8001
# ==============================